from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

doc = Document()

# Cover Page
title = doc.add_heading('SOFTWARE PROJECT BLUEPRINT (SPB)', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Tugas Ujian Tengah Semester\nMata Kuliah Proyek Pemrograman (ST165)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n\n\n[ LOGO TIM DISINI ]\n\n\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Nama Proyek: [ NAMA PROYEK ANDA ]', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('Nama Tim: [ NAMA TIM ANDA ]', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n\nAnggota Tim:\n')
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Nama Anggota'
hdr_cells[2].text = 'NIM'
hdr_cells[3].text = 'Role'

for i in range(1, 5):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = '[Nama]'
    row_cells[2].text = '[NIM]'
    row_cells[3].text = '[Role, cth: Project Manager]'

doc.add_page_break()

# Bab 1
doc.add_heading('1. Pendahuluan', level=1)

doc.add_heading('1.1 Latar Belakang', level=2)
doc.add_paragraph('Bagian ini menjelaskan identifikasi masalah yang dihadapi pengguna, dan bagaimana perangkat lunak yang akan dibangun memberikan solusi atas masalah tersebut. Jelaskan urgensi dan manfaat dari sistem ini.')

doc.add_heading('1.2 Tujuan Pengembangan Sistem', level=2)
doc.add_paragraph('Tujuan spesifik, terukur, dan relevan dari pengembangan perangkat lunak ini. Contoh:\n1. Memudahkan pelanggan dalam melakukan pemesanan secara online.\n2. Mengurangi waktu antrean di kasir fisik.')

doc.add_heading('1.3 Deskripsi Singkat Sistem', level=2)
doc.add_paragraph('Jelaskan secara ringkas ruang lingkup sistem. Apa saja yang bisa dilakukan oleh sistem, platform apa yang didukung (misal: Web, Android, iOS), dan modul utama yang tersedia.')

# Bab 2
doc.add_heading('2. Kebutuhan Sistem (System Requirements)', level=1)
doc.add_paragraph('Bagian ini diadaptasi dari standar IEEE 830 (Software Requirements Specification) untuk memastikan identifikasi kebutuhan dilakukan secara sistematis.')

doc.add_heading('2.1 Identifikasi Stakeholder & User', level=2)
table_user = doc.add_table(rows=1, cols=3)
table_user.style = 'Table Grid'
h_user = table_user.rows[0].cells
h_user[0].text = 'Kategori User'
h_user[1].text = 'Peran / Hak Akses'
h_user[2].text = 'Deskripsi'
r_user = table_user.add_row().cells
r_user[0].text = 'Admin'
r_user[1].text = 'Akses Penuh'
r_user[2].text = 'Mengelola data master dan sistem'

doc.add_heading('2.2 Functional Requirements', level=2)
table_fr = doc.add_table(rows=1, cols=3)
table_fr.style = 'Table Grid'
h_fr = table_fr.rows[0].cells
h_fr[0].text = 'ID'
h_fr[1].text = 'Nama Fitur'
h_fr[2].text = 'Deskripsi Kebutuhan'
r_fr = table_fr.add_row().cells
r_fr[0].text = 'FR-01'
r_fr[1].text = 'Login User'
r_fr[2].text = 'Sistem harus memungkinkan pengguna untuk login menggunakan email dan password.'

doc.add_heading('2.3 Use Case Diagram', level=2)
doc.add_paragraph('[ INSERT GAMBAR USE CASE DIAGRAM DISINI ]\n*Pastikan diagram jelas dan memiliki boundary, actor, serta relasi yang tepat.*')

doc.add_heading('2.4 Deskripsi Use Case (Minimal Use Case Utama)', level=2)
table_uc = doc.add_table(rows=6, cols=2)
table_uc.style = 'Table Grid'
table_uc.cell(0,0).text = 'Nama Use Case'
table_uc.cell(0,1).text = '[ Nama Use Case ]'
table_uc.cell(1,0).text = 'Aktor'
table_uc.cell(1,1).text = '[ Aktor yang terlibat ]'
table_uc.cell(2,0).text = 'Deskripsi'
table_uc.cell(2,1).text = '[ Penjelasan singkat mengenai tujuan use case ini ]'
table_uc.cell(3,0).text = 'Pre-condition'
table_uc.cell(3,1).text = '[ Status sistem sebelum use case dieksekusi ]'
table_uc.cell(4,0).text = 'Main Flow\n(Skenario Utama)'
table_uc.cell(4,1).text = '1. Aktor melakukan A\n2. Sistem merespons B\n3. ...'
table_uc.cell(5,0).text = 'Post-condition'
table_uc.cell(5,1).text = '[ Status sistem setelah use case berhasil dieksekusi ]'

doc.add_heading('2.5 Non-Functional Requirements', level=2)
doc.add_paragraph('Tentukan metrik kualitas dari sistem yang akan dibangun:')
doc.add_paragraph('1. Security: [Contoh: Data password harus di-hash menggunakan bcrypt.]')
doc.add_paragraph('2. Performance: [Contoh: Waktu respon API maksimal 2 detik.]')
doc.add_paragraph('3. Usability: [Contoh: Sistem harus dapat dioperasikan tanpa pelatihan khusus.]')

# Bab 3
doc.add_heading('3. Desain Sistem (System Design)', level=1)
doc.add_heading('3.1 Arsitektur Sistem', level=2)
doc.add_paragraph('[ Jelaskan pola arsitektur, misalnya Monolithic MVC atau Client-Server (REST API). Insert diagram arsitektur jika ada. ]')

doc.add_heading('3.2 Activity Diagram', level=2)
doc.add_paragraph('[ INSERT ACTIVITY DIAGRAM DISINI ]\n*Catatan: Hanya untuk alur use case yang kompleks/signifikan.*')

doc.add_heading('3.3 Database Design (ERD)', level=2)
doc.add_paragraph('[ INSERT GAMBAR ERD DISINI ]\n*Pastikan kardinalitas relasi (1:N, M:N) terlihat jelas.*')

doc.add_heading('3.4 High Fidelity UI Design', level=2)
doc.add_paragraph('[ INSERT SCREENSHOT DESAIN UI (FIGMA/DLL) DISINI ]')

# Bab 4
doc.add_heading('4. Repository Tim (Performance Evidence)', level=1)
doc.add_heading('4.1 Link Git Repository', level=2)
doc.add_paragraph('URL: [ https://github.com/username/repo-name ]')

doc.add_heading('4.2 Screenshot Repository', level=2)
doc.add_paragraph('[ INSERT SCREENSHOT REPOSITORY DISINI ]')

doc.add_heading('4.3 Struktur Branch', level=2)
doc.add_paragraph('Sebutkan aturan branching (branching strategy) yang disepakati:')
doc.add_paragraph('- main: Branch utama untuk rilis production.')
doc.add_paragraph('- dev: Branch untuk integrasi fitur sebelum masuk ke main.')
doc.add_paragraph('- feature/*: Branch untuk pengerjaan fitur spesifik oleh individu.')

doc.add_heading('4.4 Commit Activity', level=2)
doc.add_paragraph('[ INSERT SCREENSHOT COMMIT GRAPH / INSIGHTS CONTRIBUTORS DISINI ]')
doc.add_paragraph('Penjelasan Singkat Aktivitas Tim:')
doc.add_paragraph('[ Ceritakan bagaimana pembagian tugas terjadi. Contoh: Anggota A mengerjakan Backend di branch backend-api, Anggota B mengerjakan UI di branch feature-ui, lalu digabungkan di branch dev melalui Pull Request. ]')

doc.save('Template_UTS_Propem_Sistematis.docx')
