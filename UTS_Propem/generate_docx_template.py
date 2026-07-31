from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

doc = Document()

# Cover Page
title = doc.add_heading('SOFTWARE PROJECT BLUEPRINT (SPB)', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Tugas Ujian Tengah Semester\nMata Kuliah Proyek Pemrograman (ST165)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n\n\n[LOGO TIM DISINI]\n\n\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Nama Proyek: [NAMA PROYEK ANDA]', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('Nama Tim: [NAMA TIM ANDA]', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n\nAnggota Tim:\n')
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Nama Anggota'
hdr_cells[2].text = 'NIM'
hdr_cells[3].text = 'Role'

for i in range(1, 5):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = '[Nama Anggota]'
    row_cells[2].text = '[NIM]'
    row_cells[3].text = '[Role]'

doc.add_page_break()

# Bab 1
doc.add_heading('1. Pendahuluan', level=1)
doc.add_heading('1.1 Latar Belakang', level=2)
doc.add_paragraph('[Jelaskan permasalahan atau alasan mengapa sistem/aplikasi ini perlu dibangun]')

doc.add_heading('1.2 Tujuan Pengembangan Sistem', level=2)
doc.add_paragraph('[Sebutkan tujuan utama yang ingin dicapai dari pengembangan sistem ini]')

doc.add_heading('1.3 Deskripsi Singkat Sistem', level=2)
doc.add_paragraph('[Berikan gambaran umum mengenai sistem, fitur utamanya, dan platform apa yang digunakan (Web/Mobile/Desktop)]')

# Bab 2
doc.add_heading('2. Kebutuhan Sistem (System Requirements)', level=1)
doc.add_heading('2.1 Identifikasi Stakeholder & User', level=2)
doc.add_paragraph('[Sebutkan siapa saja stakeholder dan user yang akan menggunakan sistem ini]')

doc.add_heading('2.2 Functional Requirements (berbasis Use Case)', level=2)
doc.add_paragraph('[Sebutkan daftar kebutuhan fungsional sistem, tidak menggunakan diagram di bagian ini. Contoh: FR-01 User dapat login]')

doc.add_heading('2.3 Use Case Diagram', level=2)
doc.add_paragraph('[Insert gambar Use Case Diagram dari sistem Anda di sini]')

doc.add_heading('2.4 Deskripsi Use Case (Minimal use case utama)', level=2)
doc.add_paragraph('[Jelaskan skenario/detail alur dari use case utama. Contoh: Nama Use Case, Aktor, Pre-condition, Main Flow, Post-condition]')

doc.add_heading('2.5 Non-Functional Requirements', level=2)
doc.add_paragraph('[Sebutkan kebutuhan non-fungsional sistem. Contoh: Keamanan, Performa, Usability, dll]')

# Bab 3
doc.add_heading('3. Desain Sistem (System Design)', level=1)
doc.add_heading('3.1 Arsitektur Sistem', level=2)
doc.add_paragraph('[Jelaskan detail arsitektur yang digunakan, misalnya Client-Server, MVC. Sertakan diagram jika perlu]')

doc.add_heading('3.2 Activity Diagram', level=2)
doc.add_paragraph('[Insert Activity Diagram HANYA untuk use case yang paling kompleks atau signifikan di sini]')

doc.add_heading('3.3 Database Design (ERD)', level=2)
doc.add_paragraph('[Insert gambar Entity Relationship Diagram (ERD) di sini]')

doc.add_heading('3.4 High Fidelity UI Design', level=2)
doc.add_paragraph('[Insert screenshot desain antarmuka/UI resolusi tinggi di sini (misal dari Figma)]')

# Bab 4
doc.add_heading('4. Repository Tim (Performance Evidence)', level=1)
doc.add_heading('4.1 Link Git Repository', level=2)
doc.add_paragraph('URL: [Masukkan link repository Github/Gitlab Anda]')

doc.add_heading('4.2 Screenshot Repository', level=2)
doc.add_paragraph('[Insert screenshot tampilan utama repository]')

doc.add_heading('4.3 Struktur Branch', level=2)
doc.add_paragraph('[Sebutkan dan jelaskan branch apa saja yang dibuat dan digunakan. Contoh: main, dev, feature-x]')

doc.add_heading('4.4 Commit Activity', level=2)
doc.add_paragraph('[Insert screenshot commit activity / grafik history / ringkasan kontribusi tiap anggota dari Git]')
doc.add_paragraph('Penjelasan singkat aktivitas tim:\n[Jelaskan secara ringkas bagaimana tim berkolaborasi menggunakan git, siapa mengerjakan apa, dll]')

doc.add_page_break()
doc.add_paragraph('Catatan: Jangan lupa simpan (Export) dokumen ini ke bentuk PDF dengan ukuran maksimal 2 MB sebelum dikumpulkan ke Dashboard.')

doc.save('Template_UTS_Propem.docx')
